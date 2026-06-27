"""File processing pipeline — PDF, DOCX, images, and Google Sheets parsing."""

import base64
import io
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger("paper-assistant.file-processor")


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------


@dataclass
class Section:
    """A named section of the paper (Intro, Methods, Results, etc.)."""

    heading: str
    content: str
    start_line: int = 0
    figures: list[str] = field(default_factory=list)  # figure captions found
    references: list[str] = field(default_factory=list)  # references cited


@dataclass
class FigureInfo:
    """Metadata about an extracted figure."""

    filename: str
    caption: str = ""
    page_number: int = 0
    section: str = ""  # which section it belongs to


@dataclass
class PaperDocument:
    """Parsed paper with sections, figures, and metadata."""

    title: str = ""
    authors: str = ""
    abstract: str = ""
    sections: list[Section] = field(default_factory=list)
    figures: list[FigureInfo] = field(default_factory=list)
    full_text: str = ""
    raw_format: str = "unknown"  # "pdf", "docx", "gdoc", "text"


@dataclass
class ReviewerComment:
    """A single reviewer comment extracted from feedback."""

    id: str  # e.g., "R1-C3"
    reviewer: str  # "Reviewer 1" or "Editor"
    comment_number: int
    text: str
    severity: str = "unspecified"  # "major", "minor", "editorial"
    related_sections: list[str] = field(default_factory=list)
    related_figures: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# PDF processing
# ---------------------------------------------------------------------------


def parse_pdf(content: bytes, filename: str) -> PaperDocument:
    """Extract text and figures from a PDF."""
    import pdfplumber

    doc = PaperDocument(title=filename, raw_format="pdf")
    full_text_parts = []
    all_figures: list[FigureInfo] = []

    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()
            if text:
                full_text_parts.append(text)

            # Check for embedded images
            if hasattr(page, "images") and page.images:
                for img_idx, img in enumerate(page.images):
                    all_figures.append(
                        FigureInfo(
                            filename=f"page{i}_img{img_idx}.png",
                            page_number=i,
                        )
                    )

    doc.full_text = "\n\n".join(full_text_parts)
    doc.figures = all_figures
    doc.sections = _parse_sections(doc.full_text)
    doc.title, doc.abstract, doc.authors = _extract_metadata(doc.full_text)

    logger.info("Parsed PDF '%s': %d pages, %d sections, %d images",
                 filename, len(full_text_parts), len(doc.sections), len(all_figures))
    return doc


# ---------------------------------------------------------------------------
# DOCX processing
# ---------------------------------------------------------------------------


def parse_docx(content: bytes, filename: str) -> PaperDocument:
    """Extract text from a .docx file."""
    from docx import Document as DocxDocument

    doc = PaperDocument(title=filename, raw_format="docx")
    docx = DocxDocument(io.BytesIO(content))

    full_text_parts = []
    for para in docx.paragraphs:
        full_text_parts.append(para.text)

    doc.full_text = "\n\n".join(full_text_parts)
    doc.sections = _parse_sections(doc.full_text)
    doc.title, doc.abstract, doc.authors = _extract_metadata(doc.full_text)

    logger.info("Parsed DOCX '%s': %d paragraphs, %d sections",
                 filename, len(docx.paragraphs), len(doc.sections))
    return doc


# ---------------------------------------------------------------------------
# Markdown / plain text
# ---------------------------------------------------------------------------


def parse_text(content: str, filename: str) -> PaperDocument:
    """Parse plain text or markdown as a paper."""
    doc = PaperDocument(title=filename, raw_format="text")
    doc.full_text = content
    doc.sections = _parse_sections(content)
    doc.title, doc.abstract, doc.authors = _extract_metadata(content)

    logger.info("Parsed text '%s': %d chars, %d sections",
                 filename, len(content), len(doc.sections))
    return doc


# ---------------------------------------------------------------------------
# Section parsing
# ---------------------------------------------------------------------------


def _parse_sections(text: str) -> list[Section]:
    """Split paper text into sections based on common academic headers."""
    from config import SECTION_PATTERNS

    patterns = [re.compile(p, re.IGNORECASE | re.MULTILINE) for p in SECTION_PATTERNS]

    # Find all potential section boundaries
    boundaries: list[tuple[int, str]] = []
    for pattern in patterns:
        for match in pattern.finditer(text):
            boundaries.append((match.start(), match.group().strip()))

    boundaries.sort(key=lambda x: x[0])

    if not boundaries:
        # If no headers detected, treat the entire text as one section
        return [Section(heading="Full Text", content=text)]

    sections = []
    for i, (pos, heading) in enumerate(boundaries):
        start = pos
        end = boundaries[i + 1][0] if i + 1 < len(boundaries) else len(text)
        content = text[start:end].strip()

        # Remove the heading from content start
        if content.startswith(heading):
            content = content[len(heading):].strip()

        sections.append(Section(heading=heading, content=content, start_line=start))

    return sections


# ---------------------------------------------------------------------------
# Metadata extraction
# ---------------------------------------------------------------------------


def _extract_metadata(text: str) -> tuple[str, str, str]:
    """Heuristically extract title, abstract, and authors from the paper text."""
    title = ""
    abstract = ""
    authors = ""

    lines = text.strip().split("\n")

    # Title: first substantial line
    title_idx = -1
    for i, line in enumerate(lines[:20]):
        stripped = line.strip()
        if len(stripped) > 10 and not stripped.startswith(("#", "http", "©", "Correspondence")):
            title = stripped
            title_idx = i
            break

    # Abstract: between "Abstract" and the next section header
    abs_match = re.search(
        r"(?:^|\n)(?:#+\s*)?(?:Abstract|Summary)\s*\n+(.+?)(?:\n(?:#+\s*)?(?:Introduction|Background|Main|Results))",
        text, re.DOTALL | re.IGNORECASE
    )
    if abs_match:
        abstract = abs_match.group(1).strip()[:2000]

    # Authors: scan the few lines after the title for a plausible author list.
    # Conservative — only set when a line looks like comma/and-separated names
    # and not an affiliation or a sentence; otherwise leave empty rather than
    # risk populating garbage.
    if title_idx >= 0:
        affiliation_tokens = ("university", "department", "institute", "laborator",
                              "@", "http", "corresponding", "©", "email")
        for line in lines[title_idx + 1: title_idx + 9]:
            stripped = line.strip()
            if not (10 <= len(stripped) <= 200):
                continue
            low = stripped.lower()
            if stripped.startswith(("#", "http", "©")):
                continue
            if "," not in stripped and not re.search(r"\band\b", low):
                continue
            if stripped.endswith(".") or any(tok in low for tok in affiliation_tokens):
                continue
            # Strip affiliation superscript markers (digits, *, †, ‡, §)
            authors = re.sub(r"[\d\*†‡§]+", "", stripped)
            authors = re.sub(r",\s*,", ",", authors)
            authors = re.sub(r"\s*,\s*", ", ", authors)
            authors = re.sub(r"\s{2,}", " ", authors).strip(" ,")
            break

    return title, abstract, authors


# ---------------------------------------------------------------------------
# Reviewer comment extraction
# ---------------------------------------------------------------------------


def _is_garbled(text: str, threshold: float = 0.3) -> bool:
    """Return True if the text looks like corrupted/binary data."""
    if not text:
        return True
    # Count replacement chars (�), null bytes, and high control chars
    garbled = sum(1 for c in text if c in ('\x00', '�') or (ord(c) < 9 and c not in '\n\r\t'))
    return (garbled / max(len(text), 1)) > threshold


def _deduplicate_comments(comments: list[ReviewerComment]) -> list[ReviewerComment]:
    """Remove comments whose text is substantially similar to another."""
    seen = []
    for c in comments:
        # Compare first 100 chars against already-accepted comments
        prefix = c.text[:100].strip().lower()
        is_dup = any(
            prefix in s.text[:200].strip().lower()
            or s.text[:100].strip().lower() in prefix
            for s in seen
        )
        if not is_dup:
            seen.append(c)
    return seen


def extract_reviewer_comments(text: str) -> list[ReviewerComment]:
    """
    Extract individual reviewer comments from reviewer feedback text.

    Handles common formats:
    - "Reviewer 1, Comment 1: ..."
    - "Reviewer 1 Comments for the Author" + introductory paragraph + numbered points
    - "R1-C1: ..."
    - Numbered lists under reviewer headers
    - Editor / AE comments
    """
    # Pre-filter: strip out obviously garbled chunks
    if _is_garbled(text):
        logger.warning("Skipping entirely garbled comment file")
        return []

    comments = []
    counter = [0]  # mutable counter for unique IDs

    # --- Pattern 1: "Reviewer X, Comment Y: ..." ---
    pattern1 = re.compile(
        r"(?:Reviewer|Referee)\s*(\d+)[,:]\s*(?:Comment|Point|Issue|Question)\s*(\d+)[,:]\s*(.+?)(?=(?:Reviewer|Referee)\s*\d+|$)",
        re.DOTALL | re.IGNORECASE,
    )

    for match in pattern1.finditer(text):
        reviewer_num = match.group(1)
        comment_num = int(match.group(2))
        comment_text = match.group(3).strip()
        if _is_garbled(comment_text) or len(comment_text) < 20:
            continue
        counter[0] += 1
        comments.append(
            ReviewerComment(
                id=f"R{reviewer_num}-C{comment_num}",
                reviewer=f"Reviewer {reviewer_num}",
                comment_number=comment_num,
                text=comment_text[:3000],
                severity=_classify_severity(comment_text),
            )
        )

    # --- Pattern 2: "Reviewer X (Name)" header block with numbered points ---
    pattern2 = re.compile(
        r"(?:Reviewer|Referee)\s*(\d+)\s*(?:[:(].*?[):])?\s*\n(.+?)(?=(?:Reviewer|Referee)\s*\d+|===|Editor\b|AE\b|Associate\s+Editor|$)",
        re.DOTALL | re.IGNORECASE,
    )
    for match in pattern2.finditer(text):
        reviewer_num = match.group(1)
        block = match.group(2).strip()
        if _is_garbled(block):
            continue
        # Extract numbered points
        points = re.split(r"\n\s*(?:\d+[.)]\s*|\*\s*)", block)
        for i, point in enumerate(points, start=1):
            point = point.strip()
            if _is_garbled(point):
                continue
            if len(point) > 20:
                counter[0] += 1
                comments.append(
                    ReviewerComment(
                        id=f"R{reviewer_num}-C{i}",
                        reviewer=f"Reviewer {reviewer_num}",
                        comment_number=i,
                        text=point[:3000],
                        severity=_classify_severity(point),
                    )
                )

    # --- Pattern 3: "Reviewer X Comments for the Author" with intro paragraph + numbers ---
    pattern3 = re.compile(
        r"(?:Reviewer|Referee)\s*(\d+)\s*(?:Comments|Feedback|Report)\s*(?:for\s+the\s+Author[s]?)?[:\n]\s*(.+?)(?=(?:Reviewer|Referee)\s*\d+|Editor\b|AE\b|Associate\s+Editor|===|$)",
        re.DOTALL | re.IGNORECASE,
    )
    for match in pattern3.finditer(text):
        reviewer_num = match.group(1)
        block = match.group(2).strip()
        if _is_garbled(block):
            continue
        # Split into introductory paragraph + numbered points
        parts = re.split(r"\n\s*(?=\d+[.)]\s)", block)
        for part in parts:
            part = part.strip()
            if _is_garbled(part):
                continue
            # Identify if this is a numbered point or the intro paragraph
            num_match = re.match(r"(\d+)[.)]\s*(.+)", part, re.DOTALL)
            if num_match:
                comment_num = int(num_match.group(1))
                point_text = num_match.group(2).strip()
            else:
                # Introductory / general comment
                comment_num = 0
                point_text = part

            if len(point_text) > 20:
                counter[0] += 1
                label = f"C{comment_num}" if comment_num > 0 else "Intro"
                comments.append(
                    ReviewerComment(
                        id=f"R{reviewer_num}-{label}",
                        reviewer=f"Reviewer {reviewer_num}",
                        comment_number=comment_num if comment_num > 0 else counter[0],
                        text=point_text[:3000],
                        severity=_classify_severity(point_text),
                    )
                )

    # --- Pattern 4: Editor / AE comments ---
    editor_pattern = re.compile(
        r"(?:Editor|AE|Associate\s+Editor)\s*(?:\(.*?\))?\s*:?\s*\n?(.+?)(?=(?:Reviewer|Referee)\s*\d+|Editor\b|AE\b|Associate\s+Editor|===|$)",
        re.DOTALL | re.IGNORECASE,
    )
    for match in editor_pattern.finditer(text):
        ed_text = match.group(1).strip()
        if _is_garbled(ed_text):
            continue
        if len(ed_text) > 20:
            counter[0] += 1
            comments.append(
                ReviewerComment(
                    id=f"ED-C{counter[0]}",
                    reviewer="Editor",
                    comment_number=counter[0],
                    text=ed_text[:3000],
                    severity=_classify_severity(ed_text),
                )
            )

    # --- Fallback: free-text chunking ---
    if not comments:
        comments = _extract_from_free_text(text)

    # Deduplicate and log
    comments = _deduplicate_comments(comments)
    logger.info("Extracted %d reviewer comments from text (pre-dedup: %d)", len(comments), counter[0])
    return comments


def extract_reviewer_comments_from_sheets(
    sheets_data: dict,
) -> list[ReviewerComment]:
    """
    Extract reviewer comments from Google Sheets data.

    Expected columns (flexible order, detected by header):
    - Reviewer / Source
    - Comment / Feedback / Concern
    - Severity (optional)
    - Status (optional)
    - Response (optional)
    """
    comments = []
    for sheet_name, rows in sheets_data.items():
        if not rows:
            continue

        # Detect header row
        header = [cell.lower().strip() if cell else "" for cell in rows[0]]

        reviewer_col = _find_column(header, ["reviewer", "source", "from"])
        comment_col = _find_column(header, ["comment", "feedback", "concern", "point", "question"])
        severity_col = _find_column(header, ["severity", "priority", "type", "category"])
        response_col = _find_column(header, ["response", "reply", "answer", "draft"])

        if comment_col is None:
            # Assume first column is reviewer, second is comment
            comment_col = 1
            reviewer_col = 0

        for i, row in enumerate(rows[1:], start=1):
            if not row or len(row) <= (comment_col or 1):
                continue

            comment_text = str(row[comment_col]) if comment_col < len(row) else ""
            if not comment_text.strip() or len(comment_text.strip()) < 10:
                continue

            reviewer_name = (
                str(row[reviewer_col]) if reviewer_col is not None and reviewer_col < len(row)
                else "Unknown"
            )

            severity = "unspecified"
            if severity_col is not None and severity_col < len(row):
                severity = str(row[severity_col]).lower()

            comments.append(
                ReviewerComment(
                    id=f"{sheet_name}-C{i}",
                    reviewer=reviewer_name,
                    comment_number=i,
                    text=comment_text[:3000],
                    severity=severity,
                )
            )

    logger.info("Extracted %d reviewer comments from sheets", len(comments))
    return comments


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _find_column(header: list[str], candidates: list[str]) -> Optional[int]:
    """Find the index of a column matching any candidate name."""
    for idx, cell in enumerate(header):
        for candidate in candidates:
            if candidate in cell:
                return idx
    return None


def _classify_severity(text: str) -> str:
    """Heuristically classify comment severity."""
    text_lower = text.lower()
    major_keywords = ["major", "critical", "must", "essential", "significant flaw",
                       "fatal", "fundamental", "require"]
    minor_keywords = ["minor", "clarify", "suggestion", "could", "optional",
                       "consider", "perhaps", "might"]
    editorial_keywords = ["typo", "grammar", "spelling", "format", "reference",
                           "citation", "punctuation"]

    major_count = sum(1 for kw in major_keywords if kw in text_lower)
    minor_count = sum(1 for kw in minor_keywords if kw in text_lower)
    editorial_count = sum(1 for kw in editorial_keywords if kw in text_lower)

    if major_count > 0:
        return "major"
    elif editorial_count > minor_count:
        return "editorial"
    elif minor_count > 0:
        return "minor"
    return "unspecified"


def _extract_from_free_text(text: str) -> list[ReviewerComment]:
    """Last-resort extraction: split by blank lines or numbered items."""
    comments = []
    chunks = re.split(r"\n\s*\n", text)
    counter = 0
    for chunk in chunks:
        chunk = chunk.strip()
        if len(chunk) > 30:
            counter += 1
            comments.append(
                ReviewerComment(
                    id=f"GEN-C{counter}",
                    reviewer="General",
                    comment_number=counter,
                    text=chunk[:3000],
                    severity=_classify_severity(chunk),
                )
            )
    return comments
